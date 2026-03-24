import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re
from pypdf import PdfReader

# 1. Configuração de Interface
st.set_page_config(page_title="Fiscalização Pro: Matriz Legal Total", layout="wide", page_icon="🛡️")

# No topo do script, logo após as importações
if 'desc_detalhada' not in st.session_state:
    st.session_state['desc_detalhada'] = ""

# Função para atualizar o estado (Callback)
def update_desc():
    st.session_state.desc_detalhada = st.session_state.desc_input

st.markdown("""
    <style>
    .stCheckbox { margin-bottom: -15px; font-size: 13px; }
    .stTabs [data-baseweb="tab"] { font-weight: bold; }
    .stTextArea textarea { background-color: #f8f9fa; }
    </style>
    """, unsafe_allow_html=True)

# --- SIDEBAR: CONFIGURAÇÃO ---
st.sidebar.header("⚙️ Configuração")
api_key = st.sidebar.text_input("Google API Key", type="password")
modelo_selecionado = "gemini-1.5-pro"

if api_key:
    genai.configure(api_key=api_key)
    try:
        modelos = [m.name.replace('models/', '') for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        modelo_selecionado = st.sidebar.selectbox("Motor de IA Ativo", modelos, index=0)
    except:
        st.sidebar.error("Erro na API Key.")

# --- BASE DE DADOS CONSOLIDADAS ---

# --- 📚 MATRIZ LEGISLATIVA CONSOLIDADA (2024-2026) ---
# Referência central para fundamentação jurídica automática
QUADRO_LEGAL_REF = {
    "REN": [
        "Decreto-Lei n.º 166/2008, de 22 de Agosto (Regime Jurídico da REN)",
        "Decreto-Lei n.º 239/2012, de 2 de novembro (1.ª alteração)",
        "Decreto-Lei n.º 124/2019, de 28 de agosto (Alteração e Republicação)",
        "Decreto-Lei n.º 124/2019 com alterações do DL n.º 123/2024, de 31 de Dezembro (Vigente)",
        "Portaria n.º 419/2012, de 20 de dezembro (Condições e Requisitos Técnicos)"
    ],
    "RAN": [
        "Decreto-Lei n.º 73/2009, de 31 de Março (Regime Jurídico da RAN)",
        "Decreto-Lei n.º 199/2015, de 16 de setembro (Alteração)",
        "Portaria n.º 162/2011, de 18 de Abril (Utilizações não agrícolas)"
    ],
    "NATURA 2000": [
        "Decreto-Lei n.º 140/99, de 24 de Abril (Regime Jurídico)",
        "Decreto-Lei n.º 49/2005 de 24 de Fevereiro (Alteração)",
        "Decreto-Lei n.º 169/2001, de 25 de maio (Redação atual)"
    ],
    "ORDENAMENTO": [
        "Decreto-Lei n.º 80/2015 (RJIGT - Regime Jurídico das IG Território)",
        "Decreto-Lei n.º 555/99 (RJUE - Regime Jurídico da Urbanização e Edificação)"
    ],
    "CONTRAORDENAÇÕES": [
        "Lei n.º 50/2006, de 29 de agosto (Lei Quadro das Contraordenações Ambientais)",
        "Alterações: Lei 89/2009, Decl. Ret. 70/2009, Lei 114/2015, DL 42-A/2016, Lei 25/2019 e DL 87/2024 (Vigente)"
    ]
}

# 💧 REN - TIPOLOGIAS DETALHADAS (DL 239/2012)
ren_litoral_dict = {
    "Faixa marítima de proteção costeira": "Linha do leito até batimétrica dos 30m",
    "Praias": "Acumulação de sedimentos (areia/cascalho)",
    "Barreiras detríticas": "Restingas, barreiras soldadas e ilhas-barreira",
    "Tômbolos": "Sedimentos que ligam ilha ao continente",
    "Sapais": "Zonas intertidais com vegetação halofítica",
    "Ilhéus e rochedos emersos no mar": "Formações rochosas destacadas",
    "Dunas costeiras e dunas fósseis": "Acumulações eólicas de areia",
    "Arribas e faixas de proteção": "Vertentes abruptas e áreas adjacentes",
    "Faixa terrestre de proteção costeira": "Proteção na ausência de dunas/arribas",
    "Águas de transição": "Secções terminais sob influência salina"
}

ren_hidro_dict = {
    "Cursos de água, leitos e margens": "Terreno coberto pelas águas e faixas confinantes",
    "Lagoas e lagos": "Meios hídricos lênticos e faixas de proteção",
    "Albufeiras": "Volumes retidos por barragens para conectividade ecológica",
    "Áreas estratégicas de proteção e recarga de aquíferos": "Zonas de infiltração máxima"
}

ren_riscos_dict = {
    "Zonas adjacentes": "Risco de cheia ou ameaça do mar (ato regulamentar)",
    "Zonas ameaçadas pelo mar": "Inundações por galgamento oceânico",
    "Zonas ameaçadas pelas cheias": "Suscetíveis a transbordo de cursos de água",
    "Áreas de elevado risco de erosão hídrica": "Declive e solo propícios a perda de terra",
    "Áreas de instabilidade de vertentes": "Movimentos de massa/deslizamentos"
}
# 🚫 INTERDIÇÕES GERAIS (Artigo 20.º do DL 166/2008)
ren_interdicoes_gerais = [
    "🏗️ Operações de loteamento",
    "🧱 Obras de urbanização, construção e ampliação",
    "🛣️ Vias de comunicação e acessos",
    "🚜 Escavações e aterros (alteração da morfologia do solo)",
    "🪓 Destruição do revestimento vegetal (não agrícola/florestal)",
    "🌊 Alteração da rede de drenagem natural"
]

# 📝 REGIMES DE CONTROLO (De acordo com o DL 239/2012)
ren_regimes_controlo = [
    "🟢 Isento de procedimento (Uso compatível livre)",
    "🟡 Comunicação Prévia à CCDR (Regime regra pós-2012)",
    "🔴 Sujeito a Autorização (Casos específicos/excecionais)",
    "⭐ Relevante Interesse Público (Despacho Governamental - Art. 21.º)"
]

# 🌿 REDE NATURA 2000 - REGIÃO CENTRO (ZEC e ZPE - PSRN2000)
zec_zpe_lista = [
    "ZEC Serra da Malcata", "ZEC Serra da Estrela", "ZEC Serra de Aire e Candeeiros", 
    "ZEC Sicó/Alvaiázere", "ZEC Serra da Lousã", "ZEC Serra do Açor", 
    "ZEC Gardunha", "ZEC Rio Zêzere", "ZEC Rio Paiva", "ZEC Rio Vouga", 
    "ZEC Estuário do Mondego", "ZEC Estuário do Tejo", "ZEC Paul de Arzila", 
    "ZEC Paul do Boquilobo", "ZEC Dunas de Mira, Gândara e Gafanhas", 
    "ZEC São Pedro de Moel", "ZEC Peniche", "ZEC Berlengas", "ZEC Montejunto", 
    "ZEC Sintra-Cascais", "ZEC Arrábida/Espichel", "ZEC Complexo do Alviela", 
    "ZEC Ribeira de Alge", "ZEC Cabeço das Videiras", "ZEC Nisa/Nelas",
    "ZPE Serra da Malcata", "ZPE Serra da Estrela", "ZPE Serra da Gardunha", 
    "ZPE Beira Interior (Tejo Internacional e Erges)", "ZPE Serra da Lousã", 
    "ZPE Paul de Arzila", "ZPE Paul do Taipal", "ZPE Paul do Boquilobo", 
    "ZPE Estuário do Mondego", "ZPE Ria de Aveiro", "ZPE Estuário do Tejo"
]

# 🌿 CONDICIONANTES ART. 9.º N.º 2 (DL 140/99 - Texto Integral)
condicionantes_art9 = [
    "a) A realização de obras de construção civil fora dos perímetros urbanos, com excepção das obras de reconstrução, demolição, conservação de edifícios e ampliação desde que esta não envolva aumento de área de implantação superior a 50% da área inicial e a área total de ampliação seja inferior a 100 m2",
    "b) A alteração do uso actual do solo que abranja áreas contínuas superiores a 5 ha",
    "c) As modificações de coberto vegetal resultantes da alteração entre tipos de uso agrícola e florestal, em áreas contínuas superiores a 5 ha, considerando-se continuidade as ocupações similares que distem entre si menos de 500 m",
    "d) As alterações à morfologia do solo, com excepção das decorrentes das normais actividades agrícolas e florestais",
    "e) A alteração do uso actual dos terrenos das zonas húmidas ou marinhas, bem como as alterações à sua configuração e topografia",
    "f) A deposição de sucatas e de resíduos sólidos e líquidos",
    "g) A abertura de novas vias de comunicação, bem como o alargamento das existentes",
    "h) A instalação de infra-estruturas de electricidade e telefónicas, aéreas ou subterrâneas, de telecomunicações, de transporte de gás natural ou de outros combustíveis, de saneamento básico e de aproveitamento de energias renováveis ou similares fora dos perímetros urbanos",
    "i) A prática de actividades motorizadas organizadas e competições desportivas fora dos perímetros urbanos",
    "j) A prática de alpinismo, de escalada e de montanhismo",
    "l) A reintrodução de espécies indígenas da fauna e da flora selvagens"
]

# 🌿 ÁREAS PROTEGIDAS (RNAP)
rnap_lista = [
    "Parque Natural das Serras de Aire e Candeeiros",
    "Parque Natural da Serra da Estrela",
    "Parque Natural do Tejo Internacional",
    "Parque Natural do Douro Internacional",
    "Reserva Natural do Paul do Boquilobo",
    "Reserva Natural da Serra da Malcata",
    "Reserva Natural das Berlengas",
    "Reserva Natural do Paul de Arzila",
    "Reserva Natural das Dunas de São Jacinto",
    "Paisagem Protegida da Serra do Açor",
    "Monumento Natural do Cabo Mondego",
    "Monumento Natural das Pegadas de Dinossáurios de Ourém/Torres Novas"
]

# 🌿 ZONAMENTO (POAP / PNA / RJUE)
zonamento_tipologias = [
    "Reserva Integral", "Reserva Parcial I", "Reserva Parcial II", 
    "Proteção Parcial I", "Proteção Parcial II", "Proteção Complementar I", 
    "Proteção Complementar II", "Área de Intervenção Específica", 
    "Zona de Proteção Estrita", "Zona de Proteção de Albufeira"
]

# 🌾 RAN - MATRIZ LEGAL INTEGRAL (DL 73/2009 republicado pelo DL 199/2015)

# Transcrição Integral: Ações Interditas - Artigo 21.º
ran_interdicoes_dict = {
    "a) Operações de loteamento e obras de urbanização, construção ou ampliação, com excepção das utilizações previstas no artigo seguinte;": "Interdição de novas edificações fora das exceções legais [cite: 238, 1034]",
    "b) Lançamento ou depósito de resíduos radioactivos, resíduos sólidos urbanos, residuos industriais ou outros produtos que contenham substâncias ou microrganismos que possam alterar e deteriorar as características do solo;": "Proibição de deposição de produtos contaminantes [cite: 239, 1035]",
    "c) Aplicação de volumes excessivos de lamas nos termos da legislação aplicável, designadamente resultantes da utilização indiscriminada de processos de tratamento de efluentes;": "Violação dos limites de tratamento de efluentes no solo [cite: 240, 1036]",
    "d) Intervenções ou utilizações que provoquem a degradação do solo, nomeadamente erosão, compactação, desprendimento de terras, encharcamento, inundações, excesso de salinidade, poluição e outros efeitos perniciosos;": "Ações prejudiciais à estrutura física/química do solo [cite: 241, 1037]",
    "e) Utilização indevida de técnicas ou produtos fertilizantes e fitofarmacêuticos;": "Uso de químicos fora das normas técnicas [cite: 244, 1038]",
    "f) Deposição, abandono ou depósito de entulhos, sucatas ou quaisquer outros resíduos.": "Deposição de resíduos de construção ou veículos em fim de vida [cite: 245, 1039]"
}

# Transcrição Integral: Utilizações Permitidas - Artigo 22.º (Cruzado com Portaria 162/2011)
ran_utilizacoes_permitidas = {
    "a) Obras com finalidade agrícola, quando integradas na gestão das explorações ligadas à actividade agrícola, nomeadamente, obras de edificação, obras hidráulicas, vias de acesso, aterros e escavações, e edificações para armazenamento ou comercialização;": "Portaria 162/2011: Implantação ≤ 1% da exploração (máx 750m2). Apoios ≤ 40m2 [cite: 249, 1043, 1449]",
    "b) Construção ou ampliação de habitação para residência própria e permanente de agricultores em exploração agrícola;": "Portaria 162/2011: ATI máxima de 300m2. Requer prova de rendimento agrícola [cite: 250, 1044, 1474]",
    "c) Construção ou ampliação de habitação para residência própria e permanente dos proprietários e respectivos agregados familiares, com os limites de área e tipologia estabelecidos no regime da habitação a custos controlados...": "Portaria 162/2011: ATI máxima de 300m2 [cite: 251, 1045, 1481]",
    "d) Instalações ou equipamentos para produção de energia a partir de fontes de energia renováveis;": "Requer projeto de recuperação dos solos para parecer da DRAP [cite: 252, 1046, 1497]",
    "e) Prospecção geológica e hidrogeológica e exploração de recursos geológicos, e respectivos anexos de apoio à exploração, respeitada a legislação específica...": "Planos de lavra e PARP devem ter parecer da DRAP [cite: 253, 1048, 1516]",
    "f) Estabelecimentos industriais, comerciais ou de serviços complementares à actividade agrícola, tal como identificados no regime de licenciamento aplicável;": "Pelo menos 50% da capacidade para produtos da própria exploração [cite: 254, 1049, 1537]",
    "g) Empreendimentos de turismo no espaço rural e de turismo de habitação, bem como empreendimentos reconhecidos como turismo de natureza, complementares à actividade agrícola;": "Área de implantação total ≤ 600m2 [cite: 255, 1050, 1543]",
    "h) Instalações de recreio e lazer complementares à actividade agrícola e ao espaço rural;": "Estruturas amovíveis e necessidade justificada pela atividade [cite: 256, 1051, 1550]",
    "i) Instalações desportivas especializadas destinadas à prática de golfe, com parecer favorável pelo Turismo de Portugal, I. P., desde que não impliquem alterações irreversíveis na topografia...": "Sem alterações irreversíveis na topografia [cite: 257, 1052, 1559]",
    "j) Obras e intervenções indispensáveis à salvaguarda do património cultural, designadamente de natureza arqueológica, recuperação paisagística ou medidas de minimização...": "Determinadas pelas autoridades competentes [cite: 258, 1053, 1566]",
    "l) Obras de construção, requalificação ou beneficiação de infra-estruturas públicas rodoviárias, ferroviárias, aeroportuárias, de logística, de saneamento, de transporte e distribuição de energia eléctrica...": "Justificação da localização e medidas de minimização de ocupação RAN [cite: 259, 1054, 1572]",
    "m) Obras indispensáveis para a protecção civil;": "Sem alternativa viável fora da RAN e parecer da Proteção Civil [cite: 261, 1055, 1593]",
    "n) Obras de reconstrução e ampliação de construções já existentes, desde que estas já se destinassem e continuem a destinar-se a habitação própria;": "Portaria 162/2011: ATI total de impermeabilização ≤ 300m2 [cite: 262, 1055, 1600]",
    "o) Obras de captação de águas ou de implantação de infra-estruturas hidráulicas;": "Necessidade justificada e medidas de minimização de escavação [cite: 263, 1056, 1607]",
    "p) Obras decorrentes de exigências legais supervenientes relativas à regularização de actividades económicas previamente exercidas.": "Novo enquadramento pelo DL 199/2015 [cite: 705, 1057]"
}

# 🏛️ PATRIMÓNIO CULTURAL (Lei 107/2001)
patrimonio_interdicoes = [
    "🚫 Obra/Intervenção sem autorização da DGPC/DRC (Interior ou Exterior)",
    "🚫 Mudança de uso que afete o valor do bem classificado",
    "🚫 Destruição, danificação ou deterioração do bem",
    "🚫 Saída de bem móvel classificado do território nacional sem autorização"
]

patrimonio_condicionantes = [
    "⚠️ Intervenção em Zona Especial de Proteção (ZEP)",
    "⚠️ Intervenção em Zona de Proteção Provisória (50 metros)",
    "⚠️ Obra em imóvel em vias de classificação (Suspensão de licença)",
    "⚠️ Trabalhos arqueológicos sem autorização prévia"
]

patrimonio_deveres = [
    "❗ Incumprimento do dever de conservação e manutenção",
    "❗ Desrespeito por ordem de obras de emergência/salvaguarda",
    "❗ Violação do dever de facultar o acesso para inspeção técnica"
]

# 💧 RECURSOS HÍDRICOS (Lei n.º 58/2005 - Lei da Água)
rh_interdicoes = [
    "🚫 Utilização do Domínio Público Hídrico sem Título (Licença/Concessão)",
    "🚫 Alteração do leito ou das margens de cursos de água",
    "🚫 Extração de inertes (areia/cascalho) em locais não autorizados",
    "🚫 Descarga de águas residuais ou resíduos sem autorização (APA/ARH)",
    "🚫 Obstrução do livre fluxo das águas ou do acesso às margens"
]

rh_condicionantes = [
    "⚠️ Obras em Margem (faixa de 10m em águas não navegáveis / 50m em navegáveis)",
    "⚠️ Construções em Zonas Adjacentes (Zonas Inundáveis/Ameaçadas pelas Cheias)",
    "⚠️ Captação de águas superficiais ou subterrâneas sem balizamento/medidor",
    "⚠️ Limpeza de linhas de água com destruição de galeria ripícola autóctone"
]

# 🛠️ MEDIDAS DE MINIMIZAÇÃO E REPOSIÇÃO
medidas_minimizacao = [
    "🌱 Reposição da topografia original e do coberto vegetal autóctone",
    "🧱 Utilização de pavimentos permeáveis ou semipermeáveis",
    "🌳 Criação de cortinas arbóreas para integração paisagística",
    "💧 Implementação de sistemas de retenção e infiltração de águas pluviais",
    "🏗️ Redução da área de impermeabilização ou da cércea da edificação",
    "🚧 Remoção imediata de entulhos e resíduos de construção",
    "🛡️ Instalação de barreiras acústicas ou de contenção de poeiras"
]

# 💰 MATRIZ JURÍDICA DE SANÇÕES TOTAL (Consolidada com Património, Água e PDM)
matriz_sancionatoria = {
    "REN": {
        "Diploma": "DL 166/2008 (Art. 43.º) e Lei 50/2006 (Quadro Sancionatório Ambiental)",
        "Pessoa Singular": "2.000€ a 37.500€ (Contraordenação Muito Grave)",
        "Pessoa Coletiva": "12.000€ a 2.500.000€ (Conforme dimensão e faturação da empresa)"
    },
    "RAN": {
        "Diploma": "DL 73/2009 (Art. 39.º) republicado pelo DL 199/2015",
        "Interdições/Utilizações": "1.000€ a 3.500€ (Singular) | 1.000€ a 35.000€ (Coletiva)",
        "Deveres Acessórios": "500€ a 1.750€ (Singular) | 500€ a 17.500€ (Coletiva)",
        "Nota": "Atos administrativos que violem o regime da RAN são NULOS (Art. 38.º)"
    },
    "NATURA 2000": {
        "Diploma": "DL 140/99 (Art. 30.º) e Lei 50/2006 (Lei de Bases do Ambiente)",
        "Pessoa Singular": "2.000€ a 37.500€ (Contraordenação Muito Grave)",
        "Pessoa Coletiva": "12.000€ a 5.000.000€ (Conforme gravidade e índice de faturação)"
    },
    "PATRIMÓNIO": {
        "Diploma": "Lei n.º 107/2001 (Art. 94.º a 100.º)",
        "Pessoa Singular": "149,64€ a 3.740,98€",
        "Pessoa Coletiva": "1.496,39€ a 44.891,81€",
        "Nota": "Licenciamentos que violem normas de proteção são NULOS (Art. 5.º)"
    },
    "RECURSOS HÍDRICOS": {
        "Diploma": "Lei n.º 58/2005 (Lei da Água) e Lei n.º 50/2006",
        "Pessoa Singular": "2.000€ a 37.500€ (Infrações Muito Graves)",
        "Pessoa Coletiva": "12.000€ a 2.500.000€ (Conforme volume de negócios)"
    },
    "PDM / URBANISMO": {
        "Diploma": "DL n.º 80/2015 (RJIGT) e RJUE (DL n.º 555/99)",
        "Pessoa Singular": "500€ a 200.000€ (Dependendo da tipologia da operação)",
        "Pessoa Coletiva": "1.500€ a 450.000€ (Conforme a gravidade da infração urbanística)"
    }
}
# 🏛️ ORDENAMENTO DO TERRITÓRIO (PDM - Regime Jurídico IGT)
pdm_classes_solo = [
    "🏙️ Solo Urbano - Áreas Edificadas (Consolidadas/A expandir)",
    "🏙️ Solo Urbano - Áreas de Atividades Económicas",
    "🏙️ Solo Urbano - Espaços Verdes/Utilização Pública",
    "🌳 Solo Rústico - Espaços Agrícolas (Fora da RAN)",
    "🌲 Solo Rústico - Espaços Florestais (Produção/Conservação)",
    "🏔️ Solo Rústico - Espaços Naturais e de Proteção",
    "🏭 Solo Rústico - Áreas de Exploração de Recursos Geológicos",
    "🏚️ Solo Rústico - Aglomerados Rurais"
]
# --- INTERFACE ---
st.title("🛡️ Sistema de Fiscalização: Master Território e Ambiente")

tabs = st.tabs(["📍 Identificação", "💧 REN", "🌿 Natura & AP", "🌾 RAN", "🏛️ Património", "🌊 Recursos Hídricos", "🗺️ PDM", "📑 Informação Técnica"])

with tabs[0]:
    c1, c2 = st.columns(2)
    with c1:
        st.subheader("📍 Localização e GPS")
        local = st.text_input("Localização/Concelho", "Região Centro", key="loc_val")
        col_gps1, col_gps2 = st.columns(2)
        lat = col_gps1.text_input("Latitude", placeholder="39.xxxx", key="lat_val")
        lon = col_gps2.text_input("Longitude", placeholder="-8.xxxx", key="lon_val")
        area_m2 = st.number_input("Área Afetada (m²)", value=1000.0, key="area_val")
        
    with c2:
        st.subheader("👤 Dossier do Processo")
        inf_nome = st.text_input("Nome/Entidade", key="nome_val")
        tipo_ent = st.radio("Tipo", ["Pessoa Singular", "Pessoa Coletiva"], horizontal=True, key="tipo_val")
        
        # --- MANUTENÇÃO DOS UPLOADS ---
        upload_auto = st.file_uploader("📄 Auto de Notícia (PDF)", type=['pdf'], key="auto_noticia_pdf")
        upload_relatorio = st.file_uploader("📑 Relatório Técnico de Fiscalização (PDF)", type=['pdf'], key="relatorio_extra_pdf")
        
        st.write("---")
        # --- NOVO: INSTRUÇÕES DE AJUSTE PARA A IA ---
        instrucoes_ia = st.text_area(
            "🤖 Instruções de Ajuste para este Processo",
            placeholder="Ex: Valoriza a reincidência; foca na destruição de espécies protegidas; assume viabilidade de legalização se houver demolição parcial...",
            key="ia_custom_input",
            help="Comandos diretos que a IA deve seguir além da análise legal padrão."
        )
        
with tabs[1]:
    incide_ren = st.toggle("🚨 A infração localiza-se em área de REN?", key="switch_ren")
    
    if incide_ren:
        st.info("**Regime Jurídico da REN:** Decreto-Lei n.º 166/2008, com a redação atualizada pelo **Decreto-Lei n.º 123/2024**.")
        
        c1, c2 = st.columns([1, 1])
        with c1:
            st.subheader("1. Tipologias da REN")
            with st.expander("🌊 Áreas de Proteção do Litoral"):
                sel_litoral = st.multiselect("Subtipologias:", list(ren_litoral_dict.keys()), key="ren_litoral")
            with st.expander("💧 Ciclo Hidrológico Terrestre"):
                sel_hidro = st.multiselect("Subtipologias:", list(ren_hidro_dict.keys()), key="ren_hidro")
            with st.expander("⚠️ Prevenção de Riscos Naturais"):
                sel_riscos = st.multiselect("Subtipologias:", list(ren_riscos_dict.keys()), key="ren_riscos")
            
            sel_ren = sel_litoral + sel_hidro + sel_riscos

        with c2:
            st.subheader("2. Interdições Observadas (Art. 20.º)")
            sel_inter_ren = [i for i in ren_interdicoes_gerais if st.checkbox(i, key=f"int_ren_{i}")]
            
            st.caption("ℹ️ O Regime de Controlo e os requisitos da Portaria 419/2012 serão determinados automaticamente pela IA.")
    else:
        st.warning("Área de REN não selecionada.")
        sel_ren, sel_inter_ren = [], []

with tabs[2]:
    incide_natura = st.toggle("🌿 A infração localiza-se em Rede Natura 2000 / AP?", key="switch_natura")
    
    if incide_natura:
        st.success("**Conservação da Natureza (DL 140/99 + DL 142/2008)**")
        col1, col2 = st.columns(2)
        with col1:
            sel_zec = st.multiselect("Sítios ZEC/ZPE (Rede Natura 2000):", zec_zpe_lista)
            sel_rnap = st.multiselect("Áreas Protegidas (RNAP):", rnap_lista)
            st.write("**Condicionantes Art. 9.º n.º 2:**")
            sel_art9 = [i for i in condicionantes_art9 if st.checkbox(i, key=f"art9_{i}")]
        with col2:
            st.write("**Zonamento (POAP / PNA):**")
            sel_zon = st.multiselect("Selecione o Zonamento afetado:", zonamento_tipologias)
    else:
        st.warning("Área Natura 2000 não selecionada.")
        sel_zec, sel_rnap, sel_art9, sel_zon = [], [], [], []

with tabs[3]:
    incide_ran = st.toggle("🌾 A infração localiza-se em área de RAN?", key="switch_ran")
    
    if incide_ran:
        st.info("**Reserva Agrícola Nacional:** Decreto-Lei n.º 73/2009, republicado pelo Decreto-Lei n.º 199/2015.")
        col_r1, col_r2 = st.columns(2)
        
        with col_r1:
            st.subheader("1. Ações Interditas (Artigo 21.º)")
            # Mapeia as interdições baseadas no dicionário consolidado
            sel_inter_ran = [k for k in ran_interdicoes_dict.keys() if st.checkbox(k, key=f'ran_int_{k[:5]}')]
            
        with col_r2:
            st.subheader("2. Pretensão de Enquadramento (Artigo 22.º)")
            sel_util_ran = st.multiselect(
                "Ação enquadrada em qual alínea de utilização permitida?", 
                list(ran_utilizacoes_permitidas.keys()),
                key="util_ran_sel"
            )
            st.caption("ℹ️ A verificação de limites (Portaria 162/2011) e a viabilidade técnica serão analisadas automaticamente pela IA com base no Auto/Relatório.")

    else:
        st.warning("Área de RAN não selecionada.")
        sel_inter_ran, sel_util_ran = [], []
        
with tabs[4]:
    incide_patrimonio = st.toggle("🏛️ A infração afeta Património Cultural?", key="switch_pat")
    if incide_patrimonio:
        st.warning("**Património Cultural (Lei 107/2001)**")
        col1, col2 = st.columns(2)
        with col1:
            st.write("**Interdições e Condicionantes**")
            sel_pat_int = [i for i in patrimonio_interdicoes if st.checkbox(i, key=f'pat_int_{i}')]
            sel_pat_cond = [i for i in patrimonio_condicionantes if st.checkbox(i, key=f'pat_cond_{i}')]
        with col2:
            st.write("**Deveres e Notas**")
            sel_pat_dev = [i for i in patrimonio_deveres if st.checkbox(i, key=f'pat_dev_{i}')]
            obs_pat = st.text_area("Observações sobre o Bem Classificado:")
    else:
        st.info("Regime de Património Cultural desativado.")
        sel_pat_int, sel_pat_cond, sel_pat_dev = [], [], []

with tabs[5]:
    incide_rh = st.toggle("🌊 A infração afeta Recursos Hídricos?", key="switch_rh")
    if incide_rh:
        st.info("**Recursos Hídricos (Lei da Água - Lei n.º 58/2005)**")
        col1, col2 = st.columns(2)
        with col1:
            st.write("**Interdições**")
            sel_rh_int = [i for i in rh_interdicoes if st.checkbox(i, key=f'rh_int_{i}')]
        with col2:
            st.write("**Condicionantes**")
            sel_rh_cond = [i for i in rh_condicionantes if st.checkbox(i, key=f'rh_cond_{i}')]
    else:
        st.info("Regime de Recursos Hídricos desativado.")
        sel_rh_int, sel_rh_cond = [], []

with tabs[6]:
    incide_pdm = st.toggle("🗺️ A infração viola o PDM / Urbanismo?", key="switch_pdm")
    if incide_pdm:
        st.info("**Ordenamento do Território:** Análise baseada no RJIGT (DL 80/2015) e RJUE (DL 555/99).")
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("1. Classificação do Solo")
            sel_pdm = st.multiselect("Selecione a categoria de espaço (via Planta de Ordenamento):", pdm_classes_solo)
            artigo_pdm = st.text_input("Artigo(s) do Regulamento em análise:", placeholder="Ex: Artigo 45.º")
        
        with col2:
            st.subheader("2. Suporte Documental")
            upload_pdm = st.file_uploader("📂 Carregar Regulamento do PDM (PDF)", type=['pdf'], key="pdm_reg_upload")
            st.caption("ℹ️ A IA analisará o PDF para determinar a conformidade com os índices urbanísticos e usos permitidos.")
    
        desc_pdm = st.text_area(
            "📝 Análise Técnica de Enquadramento no PDM", 
            placeholder="Descreva a pretensão ou a violação (ex: excesso de cércea, uso interdito)...",
            height=100
        )
    else:
        st.warning("Regime de PDM desativado.")
        sel_pdm, artigo_pdm, desc_pdm = [], "", ""
        
with tabs[7]:
    st.subheader("🛠️ Medidas de Minimização Propostas")
    sel_medidas = [i for i in medidas_minimizacao if st.checkbox(i, key=f'med_{i}')]
    texto_adicional_medidas = st.text_area("Prescrições técnicas específicas:")
    
    st.divider()
    st.subheader("🏁 Finalização e Geração")
    gravidade = st.select_slider("Gravidade Proposta", options=["Leve", "Grave", "Muito Grave"])
    r_crime = st.checkbox("⚠️ Suspeita de Crime (Art. 278.º Código Penal)")
    beneficio_economico = st.checkbox("Benefício económico mensurável?")
    reincidencia = st.checkbox("Reincidência por parte do infrator?")

    st.write("---")
    st.subheader("⚖️ Regimes Sancionatórios Ativados")
    col_reg1, col_reg2 = st.columns(2)
    with col_reg1:
        # Uso de get() para evitar erro caso o tab não tenha sido aberto
        if locals().get('sel_ren'): 
            st.warning(f"🔹 **REN:** {matriz_sancionatoria['REN']['Pessoa Singular']}")
        if locals().get('sel_inter_ran'): 
            st.warning(f"🔹 **RAN:** {matriz_sancionatoria['RAN']['Interdições/Utilizações']}")
    with col_reg2:
        if locals().get('sel_zec') or locals().get('sel_art9'): 
            st.warning(f"🔹 **Natura 2000:** {matriz_sancionatoria['NATURA 2000']['Pessoa Singular']}")
        
        # CORREÇÃO DEFINITIVA: Alterado de 'AGUA' para 'RECURSOS HÍDRICOS'
        if locals().get('sel_rh_int') or locals().get('sel_rh_cond'): 
            st.warning(f"🔹 **Recursos Hídricos:** {matriz_sancionatoria['RECURSOS HÍDRICOS']['Pessoa Singular']}")

    # Função interna para exportação
    def export_docx(res_text):
        doc = Document()
        for s in doc.sections:
            s.top_margin, s.bottom_margin = Cm(2.5), Cm(2.5)
            s.left_margin, s.right_margin = Cm(3.0), Cm(2.5)
        for linha in res_text.replace('*', '').replace('#', '').split('\n'):
            linha = linha.strip()
            if not linha: continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if re.match(r'^(\d+\.|RELATÓRIO|PROPOSTA|AUTO|INFRAÇÃO|DADOS|FUNDAMENTAÇÃO|CONCLUSÃO)', linha.upper()):
                p.add_run(linha).bold = True
            else:
                p.add_run(linha)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    if st.button("🚀 Gerar Informação Técnica Fundamentada"):
        if not api_key:
            st.error("Falta a API Key no menu lateral.")
        else:
            with st.spinner("A realizar auditoria jurídica transversal e análise documental (2024-2026)..."):
                model = genai.GenerativeModel(modelo_selecionado)
                
                # --- 1. SEGURANÇA CONTRA NAMEERROR (INICIALIZAÇÃO DE VARIÁVEIS) ---
                v_ren = {'sel_ren': locals().get('sel_ren', []), 'sel_inter_ren': locals().get('sel_inter_ren', [])}
                v_ran = {'sel_inter_ran': locals().get('sel_inter_ran', []), 'sel_util_ran': locals().get('sel_util_ran', [])}
                v_natura = {'sel_zec': locals().get('sel_zec', []), 'sel_rnap': locals().get('sel_rnap', []), 'sel_art9': locals().get('sel_art9', []), 'sel_zon': locals().get('sel_zon', [])}
                v_pat = {'sel_pat_int': locals().get('sel_pat_int', []), 'sel_pat_cond': locals().get('sel_pat_cond', []), 'sel_pat_dev': locals().get('sel_pat_dev', [])}
                v_rh = {'sel_rh_int': locals().get('sel_rh_int', []), 'sel_rh_cond': locals().get('sel_rh_cond', [])}
                v_pdm = {'sel_pdm': locals().get('sel_pdm', []), 'artigo_pdm': locals().get('artigo_pdm', ''), 'desc_pdm': locals().get('desc_pdm', '')}

                # --- 2. EXTRAÇÃO DE TEXTO DOS DOCUMENTOS (AUTO, RELATÓRIO E PDM) ---
                def extrair_pdf(upload):
                    if upload:
                        try:
                            reader = PdfReader(upload)
                            return "\n".join([page.extract_text() for page in reader.pages])
                        except Exception: return "Erro na leitura do documento."
                    return ""

                texto_auto = extrair_pdf(st.session_state.get('auto_noticia_pdf'))
                texto_relatorio_extra = extrair_pdf(st.session_state.get('relatorio_extra_pdf'))
                texto_pdm_reg = extrair_pdf(st.session_state.get('pdm_reg_upload'))

                # --- 3. PREPARAÇÃO DO QUADRO LEGISLATIVO PARA A IA ---
                legis_ref_text = "\n".join([f"- {cat}: {', '.join(leis)}" for cat, leis in QUADRO_LEGAL_REF.items()])

                # --- 4. CONSTRUÇÃO DAS INSTRUÇÕES DE AJUSTE PERSONALIZADO ---
                inst_complementares = st.session_state.get('ia_custom_input', 'Seguir trâmite legal padrão.')

                instrucoes_audit = f"""
                ⚠️ INSTRUÇÃO PRIORITÁRIA DO UTILIZADOR PARA ESTE PROCESSO (SOBREPÕE-SE AO PADRÃO):
                {inst_complementares}

                ⚠️ TAREFAS DE AUDITORIA AUTOMÁTICA:
                1. REN: Determinar Regime de Controlo (Art. 21.º) e conformidade Portaria 419/2012.
                2. RAN: Analisar áreas métricas no Auto/Relatório e verificar limites Portaria 162/2011 (habitação 300m2, apoios 40m2, etc).
                3. PDM: Confrontar o Regulamento com a Análise Técnica e o Auto para ditar conformidade.
                4. SANCIONATÓRIO: Determinar GRAVIDADE e MEDIDAS DE REPOSIÇÃO via Lei 50/2006 (v. 87/2024).
                """

                # Contextos específicos para o prompt
                contexto_natura = f"NATURA 2000: Sítios: {v_natura['sel_zec']} | RNAP: {v_natura['sel_rnap']} | Art 9º: {v_natura['sel_art9']}" if incide_natura else "N/A"
                contexto_ran = f"RAN: Interditas: {v_ran['sel_inter_ran']} | Pretensão Art. 22º: {v_ran['sel_util_ran']}" if incide_ran else "N/A"

                # 5. Recuperação Segura da Descrição Manual
                desc_manual = st.session_state.get('desc_input', st.session_state.get('desc_detalhada', 'Sem descrição adicional.'))

                prompt = f"""
                Age como Perito Técnico Sénior e Jurista Forense.
                O teu objetivo é realizar uma auditoria pericial cruzada e redigir uma INFORMAÇÃO TÉCNICA FUNDAMENTADA.

                {instrucoes_audit}

                DADOS DE SUPORTE DOCUMENTAL (BASE DE PROVA):
                - AUTO DE NOTÍCIA (PDF): {texto_auto[:4000]}
                - RELATÓRIO TÉCNICO EXTRA (PDF): {texto_relatorio_extra[:4000]}
                - REGULAMENTO PDM (EXTRATO): {texto_pdm_reg[:3000]}
                
                DADOS ADICIONAIS:
                - Local: {local} | Área: {area_m2}m2 | Solo PDM: {v_pdm['sel_pdm']}
                - Descrição Manual: {desc_manual}
                - Análise Técnica PDM do Utilizador: {v_pdm['desc_pdm']}

                MATRIZ DE SERVIDÕES:
                - REN: {v_ren['sel_ren']} | Interdições: {v_ren['sel_inter_ren']}
                - {contexto_ran}
                - {contexto_natura}
                - PATRIMÓNIO: {v_pat['sel_pat_int']} | RECURSOS HÍDRICOS: {v_rh['sel_rh_int']}

                QUADRO LEGISLATIVO VIGENTE (2024-2026):
                {legis_ref_text}

                ESTRUTURA OBRIGATÓRIA:
                1. **OBJETIVO E AUDITORIA DOCUMENTAL** (Citar inconsistências entre Auto e Relatório se existirem)
                2. **ANÁLISE TÉCNICA TRANSVERSAL** (Confronto Portarias 419/2012 e 162/2011)
                3. **FUNDAMENTAÇÃO JURÍDICA E GRADUAÇÃO DA INFRAÇÃO** (Lei 50/2006 v. 2026)
                4. **PRESCRIÇÃO DE MEDIDAS DE REPOSIÇÃO E MINIMIZAÇÃO**
                5. **PARECER FINAL SOBRE VIABILIDADE DE LEGALIZAÇÃO**

                ESTILO: Jurídico, formal, Português de Portugal (PT-PT).
                """
                
                try:
                    res = model.generate_content(prompt).text
                    st.success("Informação Técnica pericial gerada com sucesso!")
                    st.write(res)
                    # Função de exportação para Word
                    st.download_button("📥 Descarregar Word", export_docx(res), file_name=f"InfoTecnica_{local}.docx")
                except Exception as e:
                    st.error(f"Erro na geração: {e}")
